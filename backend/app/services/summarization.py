"""
Document understanding and structured summarization pipeline.
"""

from __future__ import annotations

import asyncio
import re
import time
from collections import Counter
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from groq import Groq

from app.config import settings
from app.models.schemas import (
    DocumentSummaryResponse,
    SummaryInsight,
    SummaryMetric,
    SummaryMode,
)
from app.services.ingestion import get_document, get_document_file_path
from app.utils.logging import get_logger

logger = get_logger("summarization")

_summary_cache: Dict[Tuple[str, SummaryMode], DocumentSummaryResponse] = {}
_groq_client: Optional[Groq] = None

STOPWORDS = {
    "a", "about", "above", "after", "again", "against", "all", "also", "am", "an", "and",
    "any", "are", "as", "at", "be", "because", "been", "before", "being", "below", "between",
    "both", "but", "by", "can", "could", "did", "do", "does", "doing", "down", "during", "each",
    "few", "for", "from", "further", "had", "has", "have", "having", "he", "her", "here", "hers",
    "herself", "him", "himself", "his", "how", "i", "if", "in", "into", "is", "it", "its",
    "itself", "just", "me", "more", "most", "my", "myself", "no", "nor", "not", "now", "of",
    "off", "on", "once", "only", "or", "other", "our", "ours", "ourselves", "out", "over",
    "own", "same", "she", "should", "so", "some", "such", "than", "that", "the", "their",
    "theirs", "them", "themselves", "then", "there", "these", "they", "this", "those", "through",
    "to", "too", "under", "until", "up", "very", "was", "we", "were", "what", "when", "where",
    "which", "while", "who", "whom", "why", "will", "with", "would", "you", "your", "yours",
    "yourself", "yourselves",
}

HEADING_SECTION_TYPES: Dict[str, tuple[str, ...]] = {
    "abstract": ("abstract",),
    "executive_summary": ("executive summary", "summary"),
    "introduction": ("introduction", "overview", "background", "objective", "objectives"),
    "methods": ("method", "methods", "approach", "methodology", "implementation", "design", "training"),
    "results": ("results", "findings", "analysis", "evaluation", "outcomes", "observations"),
    "recommendations": ("recommendation", "recommendations", "next steps", "action items", "proposal"),
    "conclusion": ("conclusion", "conclusions", "final takeaway", "takeaway", "closing remarks"),
    "experience": ("experience", "professional experience", "work experience", "employment history"),
    "skills": ("skills", "technical skills", "core skills", "competencies"),
    "education": ("education", "certifications", "projects", "project experience"),
    "medical": ("impression", "diagnosis", "assessment", "history", "findings"),
    "legal": ("terms", "conditions", "liability", "obligations", "governing law"),
    "references": ("references", "bibliography", "appendix", "acknowledgements"),
}

SECTION_WEIGHTS: Dict[str, float] = {
    "abstract": 1.0,
    "executive_summary": 1.0,
    "results": 0.95,
    "recommendations": 0.94,
    "conclusion": 0.93,
    "introduction": 0.82,
    "methods": 0.56,
    "experience": 0.9,
    "skills": 0.86,
    "education": 0.72,
    "medical": 0.92,
    "legal": 0.74,
    "references": -0.6,
    "body": 0.66,
}

MODE_CONFIG = {
    SummaryMode.QUICK: {"insights": 3, "metrics": 3, "detail_chars": 180, "audience": "concise"},
    SummaryMode.STANDARD: {"insights": 4, "metrics": 4, "detail_chars": 240, "audience": "balanced"},
    SummaryMode.DEEP: {"insights": 6, "metrics": 6, "detail_chars": 360, "audience": "detailed"},
    SummaryMode.EXECUTIVE: {"insights": 4, "metrics": 5, "detail_chars": 230, "audience": "executive"},
    SummaryMode.STUDENT: {"insights": 4, "metrics": 4, "detail_chars": 220, "audience": "simple"},
}

LANDMARK_NOTES = {
    "attention is all you need": (
        "This paper introduced the Transformer architecture, which later became the foundation "
        "for GPT-style systems and most modern large language models."
    ),
    "general data protection regulation": (
        "This is the GDPR framework, a landmark privacy regulation that reshaped how organizations "
        "handle personal data across Europe and beyond."
    ),
    "gdpr": (
        "This is the GDPR framework, a landmark privacy regulation that reshaped how organizations "
        "handle personal data across Europe and beyond."
    ),
}

STYLE_BANNED_OPENERS = (
    "this document provides",
    "the paper focuses on",
    "it can be observed that",
    "the objective is to evaluate",
)

# ── Contact-line detection patterns (for filtering from prose fields) ──
_EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
_PHONE_RE = re.compile(r"(\+?\d[\d\-\s().]{7,}\d)")
_URL_RE = re.compile(r"https?://[^\s,)]+", re.IGNORECASE)
_SOCIAL_RE = re.compile(
    r"(?:linkedin\.com/in/|github\.com/|portfolio|behance\.net/|dribbble\.com/)[^\s,)]*",
    re.IGNORECASE,
)
_CONTACT_LABEL_RE = re.compile(
    r"^\s*(email|phone|mobile|tel|linkedin|github|portfolio|website|address|location)\s*[:\-|]",
    re.IGNORECASE,
)


def _is_contact_line(line: str) -> bool:
    """Return True if a line is primarily contact information."""
    stripped = line.strip()
    if not stripped or len(stripped) < 4:
        return False
    if _CONTACT_LABEL_RE.match(stripped):
        return True
    if _EMAIL_RE.search(stripped):
        return True
    if _SOCIAL_RE.search(stripped):
        return True
    # Phone-only line (short line with mostly digits/punctuation)
    if _PHONE_RE.search(stripped) and len(stripped) < 40:
        return True
    return False


def _strip_contact_info(text: str) -> str:
    """Remove contact information fragments from prose text."""
    result = text
    # Remove emails
    result = _EMAIL_RE.sub("", result)
    # Remove URLs
    result = _URL_RE.sub("", result)
    # Remove social handles like linkedin.com/in/... github.com/...
    result = _SOCIAL_RE.sub("", result)
    # Remove phone numbers in short fragments
    for match in _PHONE_RE.finditer(result):
        phone = match.group(0)
        if len(phone.strip()) >= 7:
            result = result.replace(phone, "")
    # Clean up leftover punctuation and whitespace
    result = re.sub(r"[|•·]\s*[|•·]", "", result)
    result = re.sub(r"\s{2,}", " ", result).strip()
    result = re.sub(r"^[,;:|•·\s]+|[,;:|•·\s]+$", "", result)
    return result


def _extract_raw_contacts(full_text: str) -> List[str]:
    """Extract raw contact info strings from document text for dedicated display.

    Deduplicates aggressively:
    - Emails: same @domain.tld → keep the longest local-part only
    - Socials: same handle (path after github.com/ or linkedin.com/in/) → keep once
    - Phones: same digit sequence → keep once
    """
    # ── Emails: keep longest local-part per @domain suffix ──
    email_by_suffix: dict[str, str] = {}
    for m in _EMAIL_RE.finditer(full_text):
        val = m.group(0)
        at = val.find("@")
        if at < 0:
            continue
        suffix = val[at:].lower()  # e.g. "@gmail.com"
        existing = email_by_suffix.get(suffix)
        if existing is None or len(val) > len(existing):
            email_by_suffix[suffix] = val

    # ── Socials: deduplicate by normalised handle ──
    _HANDLE_RE = re.compile(r"(?:linkedin\.com/in/|github\.com/)([^\s,)/?#]+)", re.IGNORECASE)
    social_by_handle: dict[str, str] = {}
    for m in _SOCIAL_RE.finditer(full_text):
        val = m.group(0).strip().rstrip("/")
        hm = _HANDLE_RE.search(val)
        handle = hm.group(1).lower() if hm else val.lower()
        if handle not in social_by_handle:
            social_by_handle[handle] = val

    # ── Phones: deduplicate by digit sequence ──
    phone_by_digits: dict[str, str] = {}
    for m in _PHONE_RE.finditer(full_text):
        val = m.group(0).strip()
        digits = re.sub(r"\D", "", val)
        if 7 <= len(digits) <= 15 and digits not in phone_by_digits:
            phone_by_digits[digits] = val

    return list(email_by_suffix.values()) + list(social_by_handle.values()) + list(phone_by_digits.values())


@dataclass
class ParsedSection:
    heading: str
    content: str
    page_start: int
    page_end: int
    section_type: str = "body"
    importance_score: float = 0.0


@dataclass
class ParsedDocument:
    document_id: str
    filename: str
    mime_type: str
    title: str
    pages: List[str]
    sections: List[ParsedSection]
    full_text: str
    document_type: str
    key_terms: List[str] = field(default_factory=list)
    landmark_note: Optional[str] = None


def invalidate_document_summary_cache(document_id: str) -> None:
    """Remove cached summaries for a document."""
    keys_to_delete = [key for key in _summary_cache if key[0] == document_id]
    for key in keys_to_delete:
        _summary_cache.pop(key, None)


async def summarize_document(
    document_id: str,
    mode: SummaryMode = SummaryMode.STANDARD,
    force_refresh: bool = False,
) -> DocumentSummaryResponse:
    """Generate a structured summary for a document."""
    cache_key = (document_id, mode)
    if not force_refresh and cache_key in _summary_cache:
        return _summary_cache[cache_key]

    metadata = get_document(document_id)
    if metadata is None:
        raise ValueError("Document not found")

    file_path = get_document_file_path(document_id)
    if file_path is None:
        raise FileNotFoundError("Uploaded file is missing from local storage")

    start = time.perf_counter()
    parsed = await asyncio.to_thread(
        _parse_document,
        document_id,
        metadata.filename,
        metadata.mime_type,
        file_path,
    )
    summary = await asyncio.to_thread(_build_summary, parsed, mode)
    _summary_cache[cache_key] = summary

    logger.info(
        "document_summary_generated",
        document_id=document_id,
        mode=mode.value,
        latency_ms=round((time.perf_counter() - start) * 1000, 2),
        confidence=summary.confidence,
        document_type=summary.document_type,
    )
    return summary


def _parse_document(
    document_id: str,
    filename: str,
    mime_type: str,
    file_path: Path,
) -> ParsedDocument:
    pages = _read_document_pages(file_path, mime_type)
    cleaned_pages = [_clean_page_text(page) for page in pages]
    full_text = "\n\n".join(page for page in cleaned_pages if page.strip()).strip()

    if len(full_text) > settings.summary_max_input_chars:
        full_text = full_text[:settings.summary_max_input_chars]

    title = _extract_title(cleaned_pages, filename)
    sections = _detect_sections(cleaned_pages, title)
    key_terms = _extract_key_terms(full_text, title)
    document_type = _classify_document(title, filename, sections, full_text)
    landmark_note = _detect_landmark_note(title, full_text)

    for section in sections:
        section.importance_score = _score_section(section, cleaned_pages, key_terms, document_type)

    return ParsedDocument(
        document_id=document_id,
        filename=filename,
        mime_type=mime_type,
        title=title,
        pages=cleaned_pages,
        sections=sorted(sections, key=lambda item: item.page_start),
        full_text=full_text,
        document_type=document_type,
        key_terms=key_terms,
        landmark_note=landmark_note,
    )


def _read_document_pages(file_path: Path, mime_type: str) -> List[str]:
    if mime_type == "application/pdf":
        return _read_pdf_pages(file_path)

    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _read_docx_pages(file_path)

    if mime_type in {"text/plain", "text/markdown"}:
        return [_safe_decode_text(file_path.read_bytes())]

    return [_safe_decode_text(file_path.read_bytes())]


def _read_pdf_pages(file_path: Path) -> List[str]:
    pages: List[str] = []
    document = fitz.open(file_path)
    try:
        for page in document:
            pages.append(page.get_text("text"))
    finally:
        document.close()
    return pages


def _read_docx_pages(file_path: Path) -> List[str]:
    document = DocxDocument(str(file_path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return ["\n".join(paragraphs)] if paragraphs else [""]


def _safe_decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1", errors="ignore")


def _clean_page_text(text: str) -> str:
    lines: List[str] = []
    for raw_line in text.replace("\x00", " ").splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if not line:
            lines.append("")
            continue
        if re.fullmatch(r"(page\s+)?\d+", line.lower()):
            continue
        lines.append(line)

    cleaned = "\n".join(lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _extract_title(pages: Sequence[str], filename: str) -> str:
    if pages:
        first_page_lines = [line.strip() for line in pages[0].splitlines() if line.strip()]
        title_parts: List[str] = []
        for line in first_page_lines[:8]:
            if _infer_section_type(line) not in {"body", "references"} and title_parts:
                break
            if re.search(r"\b(abstract|summary|introduction|executive summary)\b", line.lower()):
                break
            if len(line) < 4:
                continue
            if any(marker in line.lower() for marker in ("@", "linkedin.com", "github.com")):
                break
            title_parts.append(line)
            if len(title_parts) >= 2 or len(" ".join(title_parts)) > 120:
                break

        candidate = " ".join(title_parts).strip()
        if 6 <= len(candidate) <= 180:
            return candidate

    return Path(filename).stem.replace("_", " ").replace("-", " ").strip().title() or "Uploaded document"


def _detect_sections(pages: Sequence[str], title: str) -> List[ParsedSection]:
    sections: List[ParsedSection] = []
    current_heading = "Document opening"
    current_type = "introduction"
    current_page = 1
    current_lines: List[str] = []

    def flush(page_end: int) -> None:
        content = "\n".join(current_lines).strip()
        if content:
            sections.append(
                ParsedSection(
                    heading=current_heading,
                    content=content,
                    page_start=current_page,
                    page_end=page_end,
                    section_type=current_type,
                )
            )

    for page_index, page_text in enumerate(pages, start=1):
        lines = [line.strip() for line in page_text.splitlines()]
        for line in lines:
            if not line:
                current_lines.append("")
                continue

            if page_index == 1 and line == title:
                continue

            if _is_heading_candidate(line):
                flush(page_index)
                current_heading = line
                current_type = _infer_section_type(line)
                current_page = page_index
                current_lines = []
                continue

            current_lines.append(line)

    flush(len(pages) if pages else 1)

    merged_sections: List[ParsedSection] = []
    for section in sections:
        if merged_sections and len(section.content) < 120 and section.section_type in {"body", "introduction"}:
            previous = merged_sections[-1]
            previous.content = f"{previous.content}\n{section.content}".strip()
            previous.page_end = max(previous.page_end, section.page_end)
        else:
            merged_sections.append(section)

    if not merged_sections:
        merged_sections.append(
            ParsedSection(
                heading="Document content",
                content="\n\n".join(page for page in pages if page.strip()),
                page_start=1,
                page_end=max(len(pages), 1),
                section_type="body",
            )
        )

    return merged_sections


def _is_heading_candidate(line: str) -> bool:
    cleaned = line.strip().strip(":")
    if len(cleaned) < 3 or len(cleaned) > 120:
        return False
    if cleaned.endswith((".", "?", "!")) and len(cleaned.split()) > 6:
        return False
    if re.match(r"^[-*•]\s+", cleaned):
        return False
    if re.fullmatch(r"[0-9.\- ]+", cleaned):
        return False
    if re.match(r"^\d+(\.\d+)*\s+[A-Z]", cleaned):
        return True
    if _infer_section_type(cleaned) != "body":
        return True

    words = cleaned.split()
    if len(words) > 12:
        return False

    uppercase_words = sum(1 for word in words if word[:1].isupper())
    if cleaned.isupper() or uppercase_words >= max(1, int(len(words) * 0.7)):
        return True

    return False


def _infer_section_type(heading: str) -> str:
    normalized = re.sub(r"^\d+(\.\d+)*\s*", "", heading.lower()).strip(" :.-")
    for section_type, patterns in HEADING_SECTION_TYPES.items():
        if any(pattern in normalized for pattern in patterns):
            return section_type
    return "body"


def _extract_key_terms(full_text: str, title: str, limit: int = 12) -> List[str]:
    token_counts: Counter[str] = Counter()
    tokens = re.findall(r"[A-Za-z][A-Za-z\-]{2,}", f"{title} {full_text}".lower())
    for token in tokens:
        if token in STOPWORDS:
            continue
        if token in {"document", "section", "figure", "table", "references"}:
            continue
        token_counts[token] += 2 if token in title.lower() else 1

    return [term for term, _ in token_counts.most_common(limit)]


def _classify_document(
    title: str,
    filename: str,
    sections: Sequence[ParsedSection],
    full_text: str,
) -> str:
    lower_title = title.lower()
    lower_filename = filename.lower()
    lower_text = full_text.lower()
    headings = " ".join(section.heading.lower() for section in sections)
    combined = " ".join([lower_title, lower_filename, headings, lower_text[:12000]])
    scores: Counter[str] = Counter()

    academic_markers = sum(
        1
        for marker in ("abstract", "references", "bibliography", "doi", "et al.", "method", "introduction")
        if marker in combined
    )
    if academic_markers >= 2:
        scores["Research paper"] += 5
    elif academic_markers == 1:
        scores["Research paper"] += 2
    if re.search(r"\bwe (propose|present|introduce|evaluate|show)\b", lower_text[:3000]):
        scores["Research paper"] += 4
    if "results" in combined and academic_markers >= 1:
        scores["Research paper"] += 1

    if any(keyword in combined for keyword in ("resume", "curriculum vitae", "professional experience", "education", "skills")):
        scores["Resume"] += 5
    if re.search(r"\b(linkedin|github|portfolio|experience|certification)\b", combined):
        scores["Resume"] += 3

    if any(keyword in combined for keyword in ("agreement", "party", "shall", "hereby", "governing law", "indemnify")):
        scores["Legal agreement"] += 5

    if any(keyword in combined for keyword in ("patient", "diagnosis", "impression", "clinical", "medical", "physical exam")):
        scores["Medical report"] += 5

    if any(keyword in combined for keyword in ("proposal", "deliverables", "timeline", "budget", "scope of work")):
        scores["Business proposal"] += 5

    if any(keyword in combined for keyword in ("product requirements", "acceptance criteria", "user story", "functional requirement", "stakeholder")):
        scores["Product requirement document"] += 5

    if any(keyword in combined for keyword in ("lecture", "chapter", "unit", "topic", "notes", "study guide")):
        scores["Class notes"] += 5

    if any(keyword in combined for keyword in ("annual report", "revenue", "net income", "cash flow", "balance sheet", "fiscal year")):
        scores["Financial report"] += 7

    if any(keyword in combined for keyword in ("white paper", "case study", "analysis report", "findings")):
        scores["Report"] += 3

    if not scores:
        if "report" in combined:
            scores["Report"] += 2
        if "article" in combined:
            scores["Article"] += 2

    return scores.most_common(1)[0][0] if scores else "Document"


def _detect_landmark_note(title: str, full_text: str) -> Optional[str]:
    normalized_title = re.sub(r"[^a-z0-9]+", " ", title.lower()).strip()
    for label, note in LANDMARK_NOTES.items():
        if label in normalized_title:
            return note

    preview = re.sub(r"[^a-z0-9]+", " ", full_text[:1500].lower())
    for label, note in LANDMARK_NOTES.items():
        if label in preview:
            return note

    return None


def _score_section(
    section: ParsedSection,
    pages: Sequence[str],
    key_terms: Sequence[str],
    document_type: str,
) -> float:
    score = SECTION_WEIGHTS.get(section.section_type, SECTION_WEIGHTS["body"])
    if section.section_type == "references":
        return 0.0

    total_pages = max(len(pages), 1)
    page_ratio = section.page_start / total_pages
    if page_ratio <= 0.2:
        score += 0.08
    if section.section_type in {"conclusion", "recommendations", "results"}:
        score += 0.08
    if document_type == "Resume" and section.section_type in {"experience", "skills"}:
        score += 0.08
    if document_type == "Legal agreement" and section.section_type == "legal":
        score += 0.08

    lower_content = section.content.lower()
    term_hits = sum(1 for term in key_terms[:10] if re.search(fr"\b{re.escape(term)}\b", lower_content))
    score += min(0.18, term_hits * 0.02)

    metric_hits = len(_extract_metric_candidates(section.content))
    score += min(0.1, metric_hits * 0.02)

    if len(section.content) < 80:
        score -= 0.08

    return max(0.0, min(score, 1.0))


def _build_summary(parsed: ParsedDocument, mode: SummaryMode) -> DocumentSummaryResponse:
    if not parsed.full_text.strip():
        return DocumentSummaryResponse(
            document_id=parsed.document_id,
            document_name=parsed.filename,
            title=parsed.title,
            mode=mode,
            document_type=parsed.document_type,
            purpose="The file did not expose enough machine-readable text for a trustworthy summary.",
            overview=(
                f"{parsed.title} appears to be a {parsed.document_type.lower()}, but the uploaded file "
                "does not contain enough extractable text to summarize it reliably."
            ),
            executive_summary=[],
            main_insights=[
                SummaryInsight(
                    title="Text extraction was limited",
                    detail="The document likely needs OCR or a cleaner source file before it can be summarized accurately.",
                    supporting_pages=[],
                    importance_score=0.4,
                )
            ],
            why_it_matters="A truthful summary should not guess when the document text is unavailable or incomplete.",
            key_metrics=[],
            final_takeaway="The safest next step is to upload a text-readable version of the document.",
            landmark_note=parsed.landmark_note,
            confidence=0.18,
        )

    config = MODE_CONFIG[mode]
    purpose = _derive_purpose(parsed)
    insights = _select_main_insights(parsed, config["insights"], config["detail_chars"], mode)
    key_metrics = _select_key_metrics(parsed, config["metrics"])
    executive_summary = _compose_executive_summary(parsed, purpose, insights, key_metrics, mode)
    overview = _compose_overview(parsed, purpose, insights, mode)
    why_it_matters = _compose_why_it_matters(parsed, insights, mode, executive_summary)
    final_takeaway = _compose_final_takeaway(parsed, purpose, insights, mode, executive_summary)
    confidence = _compute_summary_confidence(parsed, insights, key_metrics)
    contact_info = _extract_raw_contacts(parsed.full_text) if parsed.document_type == "Resume" else None

    return DocumentSummaryResponse(
        document_id=parsed.document_id,
        document_name=parsed.filename,
        title=parsed.title,
        mode=mode,
        document_type=parsed.document_type,
        purpose=purpose,
        overview=overview,
        executive_summary=executive_summary,
        main_insights=insights,
        why_it_matters=why_it_matters,
        key_metrics=key_metrics,
        final_takeaway=final_takeaway,
        landmark_note=parsed.landmark_note,
        contact_info=contact_info,
        confidence=confidence,
        generated_at=datetime.utcnow(),
    )


def _derive_purpose(parsed: ParsedDocument) -> str:
    preferred_types = {"abstract", "executive_summary", "introduction", "experience", "medical", "legal"}
    candidate_sections = sorted(
        parsed.sections,
        key=lambda section: (section.section_type not in preferred_types, -section.importance_score, section.page_start),
    )

    for section in candidate_sections:
        sentence = _pick_best_sentence(section.content, parsed.key_terms)
        if sentence:
            purpose = _rewrite_purpose(sentence, parsed.document_type)
            cleaned = _strip_contact_info(purpose)
            return cleaned if cleaned.strip() else purpose

    fallback_map = {
        "Research paper": "The document was written to present a central research contribution and explain its evidence.",
        "Resume": "The document was written to present a candidate's background, skills, and fit for a role.",
        "Legal agreement": "The document was written to define obligations, rights, and terms between parties.",
        "Medical report": "The document was written to document clinical findings, interpretation, and next steps.",
        "Business proposal": "The document was written to persuade stakeholders around a plan, scope, and expected value.",
        "Product requirement document": "The document was written to align a team on product goals, scope, and requirements.",
        "Class notes": "The document was written to condense material into learnable points and reminders.",
        "Financial report": "The document was written to explain financial performance, risk, and business direction.",
    }
    return fallback_map.get(parsed.document_type, "The document was written to communicate its main ideas and supporting details.")


def _pick_best_sentence(text: str, key_terms: Sequence[str]) -> str:
    best_sentence = ""
    best_score = -1.0
    for sentence in _split_sentences(text):
        score = _score_sentence(sentence, key_terms)
        if score > best_score:
            best_sentence = sentence
            best_score = score
    return _sanitize_sentence(best_sentence)


def _split_sentences(text: str) -> List[str]:
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return []
    return [sentence.strip() for sentence in re.split(r"(?<=[.!?])\s+", normalized) if sentence.strip()]


def _score_sentence(sentence: str, key_terms: Sequence[str]) -> float:
    lower_sentence = sentence.lower()
    if _looks_like_reference_noise(lower_sentence):
        return 0.0
    if _is_contact_line(sentence):
        return 0.0

    score = 0.2
    if 40 <= len(sentence) <= 260:
        score += 0.18
    if any(verb in lower_sentence for verb in ("propose", "introduce", "show", "find", "conclude", "recommend", "improved", "reduced", "increased", "delivers", "focuses")):
        score += 0.22
    if _extract_metric_candidates(sentence):
        score += 0.15

    term_hits = sum(1 for term in key_terms[:8] if re.search(fr"\b{re.escape(term)}\b", lower_sentence))
    score += min(0.3, term_hits * 0.05)

    if lower_sentence.startswith(("this paper", "this report", "this document", "we ", "the proposal", "the resume")):
        score += 0.12

    return score


def _looks_like_reference_noise(text: str) -> bool:
    citation_markers = len(re.findall(r"\[[0-9,\s]+\]|\([A-Z][a-z]+,\s*\d{4}\)", text))
    return (
        citation_markers >= 2
        or "all rights reserved" in text
        or "copyright" in text
        or text.startswith("references")
        or text.startswith("bibliography")
    )


def _rewrite_purpose(sentence: str, document_type: str) -> str:
    sentence = sentence.rstrip(".")
    lower_sentence = sentence.lower()
    if lower_sentence.startswith("we "):
        sentence = f"The authors {sentence[3:]}"
    elif lower_sentence.startswith("this paper "):
        sentence = f"The paper {sentence[11:]}"
    elif lower_sentence.startswith("this report "):
        sentence = f"The report {sentence[12:]}"

    rewritten = sentence[0].upper() + sentence[1:] if sentence else sentence
    if not rewritten.endswith("."):
        rewritten = f"{rewritten}."

    fallback_prefixes = {
        "Resume": "The resume is meant to show",
        "Legal agreement": "The agreement is meant to define",
        "Medical report": "The report is meant to record",
        "Business proposal": "The proposal is meant to explain",
    }
    if document_type in fallback_prefixes and len(rewritten.split()) < 7:
        return f"{fallback_prefixes[document_type]} the document's most important points."
    return rewritten


def _select_main_insights(
    parsed: ParsedDocument,
    limit: int,
    detail_chars: int,
    mode: SummaryMode,
) -> List[SummaryInsight]:
    candidates: List[tuple[float, ParsedSection, str]] = []
    for section in sorted(parsed.sections, key=lambda item: item.importance_score, reverse=True):
        if section.importance_score <= 0:
            continue
        for sentence in _split_sentences(section.content):
            sentence_score = _score_sentence(sentence, parsed.key_terms)
            if sentence_score <= 0:
                continue
            combined_score = min(1.0, section.importance_score * 0.6 + sentence_score * 0.4)
            candidates.append((combined_score, section, sentence))

    insights: List[SummaryInsight] = []
    seen_signatures: List[set[str]] = []

    for score, section, sentence in sorted(candidates, key=lambda item: item[0], reverse=True):
        signature = _sentence_signature(sentence)
        if any(_signature_overlap(signature, existing) >= 0.52 for existing in seen_signatures):
            continue

        cleaned_sentence = _naturalize_sentence(_sanitize_sentence(sentence))
        detail = _strip_contact_info(_trim_text(_adapt_for_mode(cleaned_sentence, mode), detail_chars))
        if not detail.strip():
            continue
        title = _build_insight_title(section.heading, sentence)
        pages = sorted({section.page_start, section.page_end}) if section.page_end != section.page_start else [section.page_start]
        insights.append(
            SummaryInsight(
                title=title,
                detail=detail,
                supporting_pages=pages,
                importance_score=round(score, 4),
            )
        )
        seen_signatures.append(signature)

        if len(insights) >= limit:
            break

    if insights:
        return insights

    fallback_detail = _trim_text(_sanitize_sentence(parsed.full_text[:240]), detail_chars)
    return [
        SummaryInsight(
            title="Core content",
            detail=fallback_detail,
            supporting_pages=[1],
            importance_score=0.45,
        )
    ]


def _sentence_signature(sentence: str) -> set[str]:
    return {
        token
        for token in re.findall(r"[a-z][a-z\-]{2,}", sentence.lower())
        if token not in STOPWORDS
    }


def _signature_overlap(left: set[str], right: set[str]) -> float:
    if not left or not right:
        return 0.0
    return len(left & right) / max(len(left | right), 1)


def _build_insight_title(section_heading: str, sentence: str) -> str:
    cleaned_heading = re.sub(r"^\d+(\.\d+)*\s*", "", section_heading).strip(" :.-")
    if cleaned_heading and cleaned_heading.lower() not in {"document opening", "document content"} and len(cleaned_heading.split()) <= 8:
        return cleaned_heading

    clause = _sanitize_sentence(sentence).rstrip(".")
    words = clause.split()
    if len(words) <= 8:
        return clause
    return " ".join(words[:8]).rstrip(",;:") + "..."


def _select_key_metrics(parsed: ParsedDocument, limit: int) -> List[SummaryMetric]:
    metric_candidates: List[tuple[float, SummaryMetric]] = []
    seen_values: set[str] = set()

    for section in sorted(parsed.sections, key=lambda item: item.importance_score, reverse=True):
        if section.importance_score <= 0:
            continue
        for value, sentence in _extract_metric_candidates(section.content):
            if value in seen_values:
                continue
            metric_kind = _metric_kind(parsed.document_type, sentence, value)
            if metric_kind == "other":
                continue

            label = _build_metric_label(section.heading, sentence, metric_kind)
            context = _metric_context(sentence, value)
            if not context:
                continue
            metric = SummaryMetric(
                label=label,
                value=value,
                context=_trim_text(context, 96),
            )
            score = section.importance_score + (0.14 if metric_kind != "other" else 0.0)
            metric_candidates.append((score, metric))
            seen_values.add(value)
    deduped: List[SummaryMetric] = []
    seen_labels: set[str] = set()
    for _, metric in sorted(metric_candidates, key=lambda item: item[0], reverse=True):
        label_key = metric.label.lower()
        if label_key in seen_labels:
            continue
        deduped.append(metric)
        seen_labels.add(label_key)
        if len(deduped) >= limit:
            break

    minimum_by_type = {"Research paper": 2, "Resume": 1, "Business proposal": 2, "Financial report": 2, "Legal agreement": 1}
    if len(deduped) < minimum_by_type.get(parsed.document_type, 1):
        return []
    return deduped


def _extract_metric_candidates(text: str) -> List[tuple[str, str]]:
    candidates: List[tuple[str, str]] = []
    patterns = (
        r"[$€£]\s?\d[\d,]*(?:\.\d+)?(?:\s?(?:million|billion|m|bn|k))?",
        r"\b\d+(?:\.\d+)?%",
        r"\b\d+(?:\.\d+)?\s?(?:x|times|years?|months?|days?|weeks?|hours?|employees|users|customers|pages|participants|samples|layers|heads|tokens|projects|patents)\b",
    )
    for sentence in _split_sentences(text):
        if _looks_like_reference_noise(sentence.lower()):
            continue
        for pattern in patterns:
            for match in re.finditer(pattern, sentence, re.IGNORECASE):
                value = match.group(0).strip()
                candidates.append((value, sentence))
    return candidates


def _build_metric_label(section_heading: str, sentence: str, metric_kind: str) -> str:
    preferred_labels = {
        "accuracy": "Accuracy",
        "bleu": "BLEU Score",
        "training_time": "Training Time",
        "dataset_size": "Dataset Size",
        "revenue": "Revenue",
        "growth": "Growth",
        "budget": "Budget",
        "experience_years": "Experience",
        "skills_count": "Skills",
        "education_count": "Education",
    }
    if metric_kind in preferred_labels:
        return preferred_labels[metric_kind]

    cleaned_heading = re.sub(r"^\d+(\.\d+)*\s*", "", section_heading).strip(" :.-")
    if cleaned_heading and cleaned_heading.lower() not in {"document opening", "document content"}:
        return cleaned_heading

    words = [
        token.capitalize()
        for token in re.findall(r"[A-Za-z][A-Za-z\-]{2,}", sentence)
        if token.lower() not in STOPWORDS
    ]
    return " ".join(words[:4]) if words else "Key metric"


def _compose_overview(
    parsed: ParsedDocument,
    purpose: str,
    insights: Sequence[SummaryInsight],
    mode: SummaryMode,
) -> str:
    intro = {
        "Research paper": f"{parsed.title} is a research paper focused on a specific technical contribution.",
        "Resume": f"{parsed.title} is a resume that presents a candidate's professional profile.",
        "Legal agreement": f"{parsed.title} is a legal agreement that sets terms between parties.",
        "Medical report": f"{parsed.title} is a medical report describing clinical findings and interpretation.",
        "Business proposal": f"{parsed.title} is a business proposal outlining a plan and expected value.",
        "Product requirement document": f"{parsed.title} is a product requirements document that aligns goals, scope, and expected behavior.",
        "Class notes": f"{parsed.title} is a set of class notes that condenses material for study or review.",
        "Financial report": f"{parsed.title} is a financial report describing performance, metrics, and business direction.",
    }.get(parsed.document_type, f"{parsed.title} is a {parsed.document_type.lower()}.".replace("document document", "document"))

    evidence = [intro, purpose]
    if insights:
        evidence.append(insights[0].detail)
    if parsed.landmark_note and mode != SummaryMode.STUDENT:
        evidence.append(parsed.landmark_note)

    # Filter contact info from evidence
    evidence = [_strip_contact_info(e) for e in evidence if _strip_contact_info(e)]

    fallback = " ".join(evidence)
    result = _naturalize_sentence(_rewrite_with_local_model(
        instruction=_summary_rewrite_instruction("overview", mode),
        evidence=evidence,
        fallback=fallback,
        max_new_tokens=120,
    ))
    return _strip_contact_info(result)


def _compose_why_it_matters(
    parsed: ParsedDocument,
    insights: Sequence[SummaryInsight],
    mode: SummaryMode,
    executive_summary: Sequence[str],
) -> str:
    evidence: List[str] = []
    conclusion_sections = [
        section
        for section in parsed.sections
        if section.section_type in {"conclusion", "recommendations", "results", "experience", "legal", "medical"}
    ]
    for section in conclusion_sections[:2]:
        sentence = _pick_best_sentence(section.content, parsed.key_terms)
        if sentence:
            evidence.append(sentence)

    if parsed.landmark_note:
        evidence.append(parsed.landmark_note)

    if insights:
        evidence.append(insights[0].detail)
    if executive_summary:
        evidence.extend(executive_summary[-1:])

    if not evidence:
        evidence.append(_default_why_it_matters(parsed.document_type))

    # Filter contact info from evidence
    evidence = [_strip_contact_info(e) for e in evidence if _strip_contact_info(e)]
    if not evidence:
        evidence.append(_default_why_it_matters(parsed.document_type))

    fallback = _default_why_it_matters(parsed.document_type)
    if evidence:
        fallback = " ".join(_sanitize_sentence(item) for item in evidence[:2])

    rewritten = _rewrite_with_local_model(
        instruction=_summary_rewrite_instruction("why_it_matters", mode),
        evidence=evidence,
        fallback=fallback,
        max_new_tokens=96,
    )
    return _strip_contact_info(_naturalize_sentence(rewritten))


def _compose_final_takeaway(
    parsed: ParsedDocument,
    purpose: str,
    insights: Sequence[SummaryInsight],
    mode: SummaryMode,
    executive_summary: Sequence[str],
) -> str:
    evidence = [purpose]
    evidence.extend(insight.detail for insight in insights[:2])
    if executive_summary:
        evidence.extend(executive_summary[-1:])
    if parsed.landmark_note and mode in {SummaryMode.DEEP, SummaryMode.EXECUTIVE, SummaryMode.STANDARD}:
        evidence.append(parsed.landmark_note)

    # Filter contact info from evidence
    evidence = [_strip_contact_info(e) for e in evidence if _strip_contact_info(e)]

    fallback = _default_takeaway(parsed.document_type, insights[0].detail if insights else purpose)
    result = _naturalize_sentence(_rewrite_with_local_model(
        instruction=_summary_rewrite_instruction("final_takeaway", mode),
        evidence=evidence,
        fallback=fallback,
        max_new_tokens=110,
    ))
    return _strip_contact_info(result)


def _compose_executive_summary(
    parsed: ParsedDocument,
    purpose: str,
    insights: Sequence[SummaryInsight],
    metrics: Sequence[SummaryMetric],
    mode: SummaryMode,
) -> List[str]:
    top_sections = sorted(parsed.sections, key=lambda item: item.importance_score, reverse=True)
    selected_sentences: List[str] = []
    for section in top_sections:
        for sentence in _split_sentences(section.content):
            cleaned = _naturalize_sentence(_sanitize_sentence(sentence))
            if len(cleaned) < 40 or _looks_like_reference_noise(cleaned.lower()):
                continue
            if _is_contact_line(cleaned):
                continue
            if any(_signature_overlap(_sentence_signature(cleaned), _sentence_signature(existing)) > 0.5 for existing in selected_sentences):
                continue
            selected_sentences.append(cleaned)
            if len(selected_sentences) >= 8:
                break
        if len(selected_sentences) >= 8:
            break

    paragraph_count = 2 if mode == SummaryMode.QUICK else 3
    evidence: List[str] = [purpose]
    evidence.extend(selected_sentences[:4])
    if metrics:
        evidence.append("Key metrics include: " + ", ".join(f"{metric.label} {metric.value}" for metric in metrics[:3]))
    if parsed.landmark_note:
        evidence.append(parsed.landmark_note)

    # Filter contact info from evidence
    evidence = [_strip_contact_info(e) for e in evidence if _strip_contact_info(e)]

    rewritten = _rewrite_with_local_model(
        instruction=(
            "Write polished executive-summary paragraphs. Keep each paragraph distinct: "
            "1) what the document is and does, 2) strongest findings, 3) significance. "
            "No bullets. No repetition."
        ),
        evidence=evidence,
        fallback=" ".join(evidence[:4]),
        max_new_tokens=230,
    )
    normalized = _naturalize_sentence(rewritten).replace("  ", " ")
    chunks = [part.strip() for part in re.split(r"(?<=[.!?])\s+(?=[A-Z])", normalized) if part.strip()]
    paragraphs: List[str] = []
    current = []
    for sentence in chunks:
        current.append(sentence)
        if len(" ".join(current)) > 210:
            paragraphs.append(" ".join(current))
            current = []
        if len(paragraphs) >= paragraph_count:
            break
    if current and len(paragraphs) < paragraph_count:
        paragraphs.append(" ".join(current))

    deduped: List[str] = []
    for paragraph in paragraphs:
        if any(_signature_overlap(_sentence_signature(paragraph), _sentence_signature(prev)) > 0.48 for prev in deduped):
            continue
        deduped.append(paragraph)

    cleaned = [_strip_contact_info(p) for p in deduped]
    cleaned = [p for p in cleaned if p.strip()]
    return cleaned[:4]


def _summary_rewrite_instruction(section_name: str, mode: SummaryMode) -> str:
    audience = MODE_CONFIG[mode]["audience"]

    if mode == SummaryMode.QUICK:
        length_guide = "Write 2-3 short, direct sentences. No filler."
    elif mode == SummaryMode.DEEP:
        length_guide = "Write a thorough, detailed paragraph covering nuances and evidence."
    elif mode == SummaryMode.EXECUTIVE:
        length_guide = "Write in business language. Emphasize metrics, outcomes, and strategic impact."
    elif mode == SummaryMode.STUDENT:
        length_guide = "Write in plain, simple language a student can understand. Avoid jargon."
    else:
        length_guide = "Write a balanced, moderately detailed paragraph."

    if section_name == "overview":
        return (
            f"Write a factual document overview using only the evidence. "
            f"Tone: {audience}. Mention what the document is and what it is trying to do. "
            f"{length_guide} "
            f"Never include contact information like emails, phone numbers, or URLs."
        )
    if section_name == "why_it_matters":
        return (
            f"Explain why the document matters using only the evidence. "
            f"Tone: {audience}. Focus on significance, not filler. "
            f"{length_guide} "
            f"Never include contact information like emails, phone numbers, or URLs."
        )
    return (
        f"Write a strong concluding takeaway using only the evidence. "
        f"Tone: {audience}. Keep it truthful and concrete. "
        f"{length_guide} "
        f"Never include contact information like emails, phone numbers, or URLs."
    )


def _rewrite_with_local_model(
    instruction: str,
    evidence: Sequence[str],
    fallback: str,
    max_new_tokens: int = 150,
) -> str:
    """Rewrite text using Groq LLM for polished prose."""
    cleaned_evidence = [_sanitize_sentence(item) for item in evidence if item and item.strip()]
    if not cleaned_evidence:
        return _sanitize_sentence(fallback)

    if not settings.groq_api_key:
        return _sanitize_sentence(fallback)

    global _groq_client
    if _groq_client is None:
        _groq_client = Groq(api_key=settings.groq_api_key)

    facts_block = "\n".join(f"- {item}" for item in cleaned_evidence[:6])
    prompt = (
        f"{instruction}\n"
        "Use only the facts below. Do not add facts that are not present.\n"
        "Do not include any preamble or labels like 'Answer:' or 'Overview:'.\n"
        "Facts:\n"
        f"{facts_block}\n"
    )

    try:
        response = _groq_client.chat.completions.create(
            model=settings.groq_model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a precise document summarizer. Write exactly what is asked, "
                        "using only the provided facts. Be concise and factual. "
                        "Never start with 'This document' or 'The document'."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=max_new_tokens * 2,
            temperature=0.3,
        )
        text = (response.choices[0].message.content or "").strip()
        cleaned = _sanitize_generated_text(text)
        return cleaned if len(cleaned) >= 30 else _sanitize_sentence(fallback)
    except Exception as exc:
        logger.warning("summary_rewrite_failed", error=str(exc))
        return _sanitize_sentence(fallback)


def _sanitize_generated_text(text: str) -> str:
    cleaned = re.sub(r"^(answer|summary|overview|final takeaway)\s*:\s*", "", text.strip(), flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+", " ", cleaned)
    if cleaned and not cleaned.endswith("."):
        cleaned = f"{cleaned}."
    return _naturalize_sentence(cleaned)


def _naturalize_sentence(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", text or "").strip()
    if not cleaned:
        return ""

    lower = cleaned.lower()
    for opener in STYLE_BANNED_OPENERS:
        if lower.startswith(opener):
            cleaned = re.sub(r"^[^.?!]*\b(?:provides|focuses on|is to evaluate|can be observed that)\b\s*", "", cleaned, flags=re.IGNORECASE).strip()
            cleaned = cleaned[0].upper() + cleaned[1:] if cleaned else ""
            break

    replacements = (
        (r"\bthis paper\b", "the paper"),
        (r"\bthis document\b", "the document"),
    )
    for pattern, replacement in replacements:
        cleaned = re.sub(pattern, replacement, cleaned, flags=re.IGNORECASE)

    return cleaned


def _metric_kind(document_type: str, sentence: str, value: str) -> str:
    lower = sentence.lower()
    if document_type == "Research paper":
        if "bleu" in lower:
            return "bleu"
        if "accuracy" in lower or "%" in value:
            return "accuracy"
        if any(token in lower for token in ("hours", "days", "steps", "epochs", "training")):
            return "training_time"
        if any(token in lower for token in ("dataset", "samples", "tokens", "vocabulary", "corpus")):
            return "dataset_size"
    if document_type in {"Business proposal", "Financial report"}:
        if any(token in lower for token in ("revenue", "sales", "income")):
            return "revenue"
        if any(token in lower for token in ("growth", "cagr", "increase", "decrease", "%")):
            return "growth"
        if any(token in lower for token in ("budget", "cost", "spend", "expense")):
            return "budget"
    if document_type == "Resume":
        if any(token in lower for token in ("year", "years")):
            return "experience_years"
        if any(token in lower for token in ("skills", "stack", "technologies")):
            return "skills_count"
        if any(token in lower for token in ("degree", "education", "university", "gpa")):
            return "education_count"
        if "%" in value:
            return "growth"
        if any(token in lower for token in ("users", "customers", "clients", "projects")):
            return "experience_years"
    if document_type == "Legal agreement":
        if any(token in lower for token in ("month", "months", "year", "years", "days", "weeks")):
            return "training_time"
        if "$" in value or "€" in value or "£" in value:
            return "budget"
    return "other"


def _metric_context(sentence: str, value: str) -> str:
    cleaned = _naturalize_sentence(_sanitize_sentence(sentence))
    if cleaned.lower().startswith(value.lower()):
        return cleaned
    return cleaned


def _default_why_it_matters(document_type: str) -> str:
    return {
        "Research paper": "It matters because it clarifies the document's core contribution and the evidence behind it.",
        "Resume": "It matters because it helps a reviewer quickly judge role fit, strengths, and professional trajectory.",
        "Legal agreement": "It matters because it determines the obligations, risk, and protections each party is accepting.",
        "Medical report": "It matters because it affects clinical understanding, follow-up, and decision-making.",
        "Business proposal": "It matters because it connects the proposed work to expected value, tradeoffs, and execution decisions.",
        "Product requirement document": "It matters because it aligns teams on scope, priorities, and the outcomes the product should deliver.",
        "Class notes": "It matters because it highlights the concepts most worth remembering and reviewing.",
        "Financial report": "It matters because it reveals business performance, risk, and the numbers that shape decisions.",
    }.get(document_type, "It matters because it distills the document into the points a reader actually needs to keep.")


def _default_takeaway(document_type: str, leading_point: str) -> str:
    return {
        "Research paper": f"The clearest takeaway is that the document's main contribution is {leading_point.lower()}",
        "Resume": f"The clearest takeaway is that the candidate is strongest where the resume shows {leading_point.lower()}",
        "Legal agreement": f"The clearest takeaway is that the agreement is primarily about {leading_point.lower()}",
        "Medical report": f"The clearest takeaway is that the report centers on {leading_point.lower()}",
        "Business proposal": f"The clearest takeaway is that the proposal's value rests on {leading_point.lower()}",
        "Product requirement document": f"The clearest takeaway is that the product effort should stay focused on {leading_point.lower()}",
        "Class notes": f"The clearest takeaway is that the material keeps returning to {leading_point.lower()}",
        "Financial report": f"The clearest takeaway is that business performance should be read through {leading_point.lower()}",
    }.get(document_type, f"The clearest takeaway is that the document is mainly about {leading_point.lower()}") + "."


def _adapt_for_mode(text: str, mode: SummaryMode) -> str:
    if mode == SummaryMode.STUDENT:
        replacements = (
            ("therefore", "so"),
            ("however", "but"),
            ("approximately", "about"),
            ("utilize", "use"),
        )
        for source, target in replacements:
            text = re.sub(fr"\b{source}\b", target, text, flags=re.IGNORECASE)
    return text


def _trim_text(text: str, limit: int) -> str:
    text = text.strip()
    if len(text) <= limit:
        return text
    trimmed = text[:limit].rsplit(" ", 1)[0].rstrip(",;:")
    return f"{trimmed}..."


def _sanitize_sentence(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", text or "").strip()
    cleaned = re.sub(r"^[•\-*]\s*", "", cleaned)
    return cleaned


def _compute_summary_confidence(
    parsed: ParsedDocument,
    insights: Sequence[SummaryInsight],
    metrics: Sequence[SummaryMetric],
) -> float:
    score = 0.25
    if parsed.title and parsed.title.lower() != Path(parsed.filename).stem.lower():
        score += 0.1
    score += min(0.2, len(parsed.full_text) / 6000)
    score += min(0.15, len(parsed.sections) * 0.02)
    score += min(0.15, len(insights) * 0.03)
    score += min(0.08, len(metrics) * 0.02)
    top_section_score = max((section.importance_score for section in parsed.sections), default=0.0)
    score += top_section_score * 0.18
    if parsed.landmark_note:
        score += 0.04
    return round(min(score, 0.98), 4)
